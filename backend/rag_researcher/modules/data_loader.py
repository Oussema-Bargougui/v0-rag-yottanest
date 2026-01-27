#!/usr/bin/env python3
"""
RAG-First Production Data Loader for AML/KYC RAG Pipeline

This module handles loading and preprocessing of various document types
for enterprise-grade RAG pipeline with RAG-FIRST design.

Key Features:
- Extracts clean, compact, linear text only (no layout blocks, no bbox)
- Image handling: inline injection with captions and LLM descriptions
- Table handling: real tables only, converted to inline text
- OCR fallback for scanned PDFs (Tesseract)
- Page-level organization with clean merged text
- Bank-grade error handling and logging
- SEMANTIC DETECTION to prevent misclassification
- Image hash caching to avoid duplicate LLM calls

Output is RAG-ready: clean text that improves retrieval accuracy.

Author: Yottanest Team
Version: 5.1.0 - Production with LLM Image Understanding
"""

import os
import uuid
import logging
import base64
import json
import tempfile
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, Tuple
from datetime import datetime
import requests

# Import configuration
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import Config

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Import document processing libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not available. Install with: pip install PyMuPDF")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Install with: pip install python-docx")

try:
    from PIL import Image
    import io
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("Pillow not available. Install with: pip install Pillow")

# Import professional table extraction libraries
try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False
    logger.warning("Camelot not available. Install with: pip install camelot-py[cv]")

try:
    import tabula
    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False
    logger.warning("Tabula not available. Install with: pip install tabula-py")

# Import OCR libraries
try:
    import pytesseract
    import pdf2image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logger.warning("Tesseract not available. Install with: pip install pytesseract pdf2image")


class MultimodalDataLoader:
    """
    RAG-first production data loader for enterprise AML/KYC systems.
    
    This class provides document processing with RAG-optimized output:
    - Clean, linear text per page (no blocks, no bbox)
    - Inline image injections with LLM-generated captions and descriptions
    - Inline table text representations
    - OCR fallback for scanned documents
    - Image hash caching to avoid duplicate LLM calls
    """
    
    # Configuration constants
    MIN_TEXT_THRESHOLD = 50  # Minimum text length to consider page not scanned
    OCR_TEXT_THRESHOLD = 100  # Minimum text length for OCR-processed pages
    MAX_TEXT_FOR_TABLE = 1500  # Max text length to consider page as table (semantic filter)
    MIN_IMAGE_SIZE = 100  # Min width/height for image processing (px)
    MIN_IMAGE_SIZE_FOR_LLM = 100  # Lowered threshold - process ALL images >=100px
    LLM_TIMEOUT = 45  # Increased timeout for LLM API calls (seconds)
    MAX_LLM_RETRIES = 2  # Increased retries (max 3 attempts total)
    
    def __init__(self):
        """Initialize RAG-first data loader."""
        # Ensure storage directories exist
        self.storage_path = Config.get_storage_path()
        self.images_path = self.storage_path / "images"
        self.images_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize image cache
        self.image_cache = self._load_image_cache()
        
        logger.info("MultimodalDataLoader initialized (RAG-First Production v5.1)")
        logger.info(f"Storage path: {self.storage_path}")
        logger.info(f"Images path: {self.images_path}")
        logger.info(f"Image cache entries: {len(self.image_cache)}")
        
        # Log available extraction methods
        self._log_extraction_methods()
    
    def _log_extraction_methods(self):
        """Log available table extraction methods."""
        methods = []
        if CAMELOT_AVAILABLE:
            methods.append("Camelot (lattice + stream)")
        if TABULA_AVAILABLE:
            methods.append("Tabula")
        if PYMUPDF_AVAILABLE:
            methods.append("PyMuPDF (fallback)")
        if TESSERACT_AVAILABLE:
            methods.append("Tesseract OCR")
        
        logger.info(f"Available table extraction methods: {', '.join(methods)}")
    
    def _load_image_cache(self) -> Dict[str, Dict[str, str]]:
        """Load image cache from disk."""
        cache_file = self.storage_path / "image_cache.json"
        if cache_file.exists():
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Failed to load image cache: {str(e)}")
        return {}
    
    def _save_image_cache(self):
        """Save image cache to disk."""
        cache_file = self.storage_path / "image_cache.json"
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.image_cache, f, indent=2)
            logger.debug(f"Saved image cache with {len(self.image_cache)} entries")
        except Exception as e:
            logger.warning(f"Failed to save image cache: {str(e)}")
    
    def _get_image_hash(self, image_data: bytes) -> str:
        """Calculate SHA256 hash of image data."""
        return hashlib.sha256(image_data).hexdigest()
    
    def process_uploaded_file(self, 
                         content: bytes, 
                         filename: str,
                         file_type: str = 'pdf') -> Optional[Dict[str, Any]]:
        """
        Process an uploaded file and extract RAG-ready content.
        
        Args:
            content: File content as bytes
            filename: Original filename
            file_type: File type ('pdf', 'docx', 'txt', 'md')
            
        Returns:
            RAG-ready document dictionary with clean text, pages, and full_text
        """
        try:
            # Sanitize filename to prevent WindowsPath crash
            try:
                filename = str(filename)
                filename = filename.split(";")[0].strip()
                logger.info(f"Processing uploaded file: {filename} ({len(content)} bytes)")
            except Exception as e:
                logger.error(f"Filename sanitization failed: {str(e)}")
                filename = "unknown_file"
            
            # Generate document ID
            doc_id = str(uuid.uuid4())
            
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_type}') as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name
            
            try:
                # Process based on file type
                if file_type == 'pdf':
                    result = self._process_pdf(tmp_path, filename, doc_id)
                elif file_type == 'docx':
                    result = self._process_docx(tmp_path, filename, doc_id)
                elif file_type in ['txt', 'md', 'text']:
                    result = self._process_text_file(tmp_path, filename, doc_id, file_type)
                else:
                    raise ValueError(f"Unsupported file type: {file_type}")
                
                # Add metadata
                result['metadata'] = {
                    'source': 'upload',
                    'file_type': file_type,
                    'created_at': datetime.now().isoformat(),
                    'original_filename': filename,
                    'file_size': len(content)
                }
                
                # Calculate full_text (all pages joined)
                full_text = '\n\n'.join(page['text'] for page in result.get('pages', []))
                result['full_text'] = full_text
                
                logger.info(f"Successfully processed {filename}: {len(result.get('pages', []))} pages, {len(full_text)} characters")
                
                # Save cache after processing
                self._save_image_cache()
                
                return result
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            return None
    
    def _process_pdf(self, file_path: str, filename: str, doc_id: str) -> Dict[str, Any]:
        """Process PDF file with RAG-first extraction."""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF is required for PDF processing")
        
        document = {
            'doc_id': doc_id,
            'filename': filename,
            'pages': []
        }
        
        pdf_document = None
        try:
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(file_path)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Extract clean text (no blocks, no bbox)
                text_content = page.get_text()
                
                # Check if page is scanned (low text content)
                is_scanned = len(text_content.strip()) < self.MIN_TEXT_THRESHOLD
                
                if is_scanned and TESSERACT_AVAILABLE:
                    logger.info(f"Page {page_num + 1} appears to be scanned, using OCR")
                    # Use OCR for scanned pages
                    ocr_text = self._extract_text_with_ocr(file_path, page_num)
                    clean_text = self._clean_and_merge_text(ocr_text) if ocr_text else ""
                    page_metadata = {'ocr_processed': True}
                else:
                    # Clean and merge text (no blocks)
                    clean_text = self._clean_and_merge_text(text_content)
                    page_metadata = {'ocr_processed': False}
                
                # SEMANTIC FILTER: Check if page contains real tables
                if self._should_skip_table_extraction(page, text_content):
                    logger.info(f"Page {page_num + 1}: Skipping table extraction (semantic filter)")
                    table_text = ""
                else:
                    # Extract tables and convert to inline text
                    table_text = self._extract_tables_as_inline_text(page, page_num + 1, file_path)
                
                # Extract images and generate inline text
                image_texts, image_metadata = self._extract_images_as_inline_text(page, page_num + 1, doc_id, filename)
                
                # Merge all content: text + tables + images
                if table_text:
                    clean_text = f"{clean_text}\n\n{table_text}"
                
                for img_text in image_texts:
                    if img_text:
                        clean_text = f"{clean_text}\n\n{img_text}"
                
                # Create page object (RAG-first format)
                page_data = {
                    'page_number': page_num + 1,
                    'text': clean_text.strip(),
                    'metadata': {
                        'has_images': len(image_texts) > 0,
                        'image_count': image_metadata.get('image_count', 0),
                        'image_ids': image_metadata.get('image_ids', []),
                        'has_tables': bool(table_text),
                        'ocr_processed': page_metadata['ocr_processed']
                    }
                }
                
                document['pages'].append(page_data)
            
        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            # Return partial document on error (fail-safe)
            if document['pages']:
                logger.warning(f"Returning partial document with {len(document['pages'])} pages")
                return document
            raise Exception(f"PDF processing failed: {str(e)}")
        finally:
            # Always close PDF document
            if pdf_document:
                try:
                    pdf_document.close()
                except:
                    pass
        
        return document
    
    def _clean_and_merge_text(self, text: str) -> str:
        """
        Clean and merge text into linear format for RAG.
        
        Operations:
        - Remove extra newlines
        - Fix hyphenation
        - Collapse whitespace
        - Preserve reading order
        - NO layout blocks, NO bbox, NO font info
        
        Args:
            text: Raw text content
            
        Returns:
            Clean, merged text string
        """
        if not text:
            return ""
        
        # Remove hyphenation at line breaks
        text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Split into lines
        lines = text.split('\n')
        
        # Merge lines that are clearly part of same paragraph
        merged_lines = []
        buffer = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if buffer:
                    merged_lines.append(' '.join(buffer))
                    buffer = []
                merged_lines.append("")  # Preserve paragraph breaks
            else:
                if buffer and not buffer[-1].endswith(('.', '!', '?', ':')):
                    # Continue same paragraph
                    buffer.append(line)
                else:
                    if buffer:
                        merged_lines.append(' '.join(buffer))
                    buffer = [line]
        
        if buffer:
            merged_lines.append(' '.join(buffer))
        
        # Rebuild text with proper spacing
        result = '\n'.join(merged_lines)
        
        # Collapse multiple spaces
        result = re.sub(r' +', ' ', result)
        
        # Remove multiple empty lines
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        return result.strip()
    
    def _should_skip_table_extraction(self, page, text_content: str) -> bool:
        """
        SEMANTIC FILTER: Check if page should skip table extraction entirely.
        """
        text_len = len(text_content.strip())
        
        # SKIP CONDITION: Page is too long (article/paragraph text)
        if text_len > self.MAX_TEXT_FOR_TABLE and text_content.count("\n") < 20:
            logger.info(f"Semantic filter: Page has {text_len} chars - too long for table")
            return True
        
        # Check for table-like characteristics
        lines = text_content.split('\n')
        lines = [line.strip() for line in lines if line.strip()]
        
        if not lines:
            return False
        
        # Check for multiple columns
        has_multiple_columns = False
        for line in lines[:10]:
            if '\t' in line or '  ' in line:
                has_multiple_columns = True
                break
        
        if not has_multiple_columns:
            logger.info(f"Semantic filter: No multiple columns detected")
            return True  # Skip - no table structure
        
        # Check for numeric data
        numeric_count = 0
        for line in lines[:10]:
            if re.search(r'\d', line):
                numeric_count += 1
        
        if numeric_count < 2:
            logger.info(f"Semantic filter: Insufficient numeric data")
            return True  # Skip - insufficient numeric data
        
        # All checks passed - this page likely contains a table
        logger.info(f"Semantic filter: PASSED - page likely contains table")
        return False
    
    def _hard_filter_fake_tables(self, tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        HARD FILTER: Remove fake tables after extraction.
        """
        filtered_tables = []
        
        for table in tables:
            headers = table.get('headers', [])
            rows = table.get('rows', [])
            table_id = table.get('table_id', 'unknown')
            
            # FILTER: Empty headers
            if not headers or not any(h.strip() for h in headers):
                logger.info(f"Hard filter: Discarding table {table_id} - empty headers")
                continue
            
            # FILTER: Insufficient rows
            if len(rows) < 2:
                logger.info(f"Hard filter: Discarding table {table_id} - only {len(rows)} rows")
                continue
            
            # FILTER: Numeric ratio < 30%
            numeric_cells = 0
            total_cells = 0
            for row in rows:
                for cell in row:
                    if cell.strip():
                        total_cells += 1
                        if re.search(r'[\d,]+\.?\d*', cell.strip()):
                            numeric_cells += 1
            
            if total_cells > 0:
                numeric_ratio = numeric_cells / total_cells
                if numeric_ratio < 0.3:
                    logger.info(f"Hard filter: Discarding table {table_id} - numeric ratio {numeric_ratio:.1%} < 30%")
                    continue
            
            # Table passed all filters
            filtered_tables.append(table)
        
        return filtered_tables
    
    def _extract_tables_as_inline_text(self, page, page_num: int, file_path: str) -> str:
        """
        Extract tables and convert to inline text for RAG.
        """
        tables = []
        
        try:
            # Strategy 1: Camelot (lattice)
            if CAMELOT_AVAILABLE:
                tables = self._extract_with_camelot(file_path, page_num - 1, 'lattice')
            
            # Strategy 2: Camelot (stream)
            if CAMELOT_AVAILABLE and not tables:
                tables = self._extract_with_camelot(file_path, page_num - 1, 'stream')
            
            # Strategy 3: Tabula
            if TABULA_AVAILABLE and not tables:
                tables = self._extract_with_tabula(file_path, page_num - 1)
            
            # Strategy 4: PyMuPDF fallback
            if PYMUPDF_AVAILABLE and not tables:
                tables = self._extract_with_pymupdf(page, page_num)
            
            # HARD FILTER: Remove fake tables
            if tables:
                filtered_tables = self._hard_filter_fake_tables(tables)
                if filtered_tables:
                    # Convert tables to inline text
                    table_texts = [self._table_to_text(table) for table in filtered_tables]
                    return '\n\n'.join(table_texts)
            
            return ""
            
        except Exception as e:
            logger.error(f"Table extraction failed for page {page_num}: {str(e)}")
            return ""
    
    def _normalize_table_to_semantic_text(self, table: Dict[str, Any]) -> str:
        """
        Normalize table and convert to semantic text for RAG.
        
        STAGE 2 TABLE NORMALIZATION:
        1. Clean table (remove footnotes, repeated headers, paragraph text)
        2. Detect header row
        3. Keep only numeric rows
        4. Generate heuristic summary (NO LLM)
        5. Convert to semantic text format
        
        Args:
            table: Dictionary with 'headers', 'rows', 'table_id'
            
        Returns:
            Semantic text block with title, summary, and data table
        """
        try:
            headers = table.get('headers', [])
            rows = table.get('rows', [])
            table_id = table.get('table_id', 'unknown')
            
            if not rows:
                return ""
            
            # STEP 1: Clean table - remove footnotes and non-numeric rows
            cleaned_rows = []
            for row in rows:
                # Check if row has numeric data
                has_numeric = any(re.search(r'[\d,]+\.?\d*', cell) for cell in row)
                
                # Check if it's a header row (repeated)
                is_header = False
                if headers:
                    # Check if row matches headers
                    header_match_count = sum(1 for h, r in zip(headers, row) 
                                           if h and r and h.lower() in r.lower())
                    if header_match_count >= min(2, len(headers)):
                        is_header = True
                
                # Check if it's a footnote/paragraph (lots of text, no numbers)
                is_footnote = False
                if not has_numeric:
                    # If row has long text, it's probably a footnote
                    text_length = sum(len(cell) for cell in row)
                    if text_length > 100:
                        is_footnote = True
                
                # Keep only numeric rows that aren't headers or footnotes
                if has_numeric and not is_header and not is_footnote:
                    cleaned_rows.append(row)
                elif is_header:
                    logger.debug(f"Removing repeated header row from table {table_id}")
                elif is_footnote:
                    logger.debug(f"Removing footnote row from table {table_id}")
            
            if not cleaned_rows:
                logger.warning(f"No valid data rows after cleaning table {table_id}")
                return ""
            
            # STEP 2: Generate heuristic summary (NO LLM)
            summary = self._generate_heuristic_table_summary(headers, cleaned_rows)
            
            # STEP 3: Build semantic text block
            lines = []
            lines.append("[TABLE]")
            
            # Title (auto or inferred)
            if headers:
                title = headers[0] if len(headers) == 1 else f"Data table ({len(cleaned_rows)} rows)"
                lines.append(f"Title: {title}")
            
            # Summary
            lines.append("Summary:")
            lines.append(summary)
            
            # Data section
            lines.append("Data:")
            
            # Headers row
            if headers:
                header_line = " | ".join(str(h) if h else "" for h in headers)
                lines.append(f"| {header_line} |")
            
            # Data rows (preserve numbers exactly)
            for row in cleaned_rows:
                row_line = " | ".join(str(cell) if cell else "" for cell in row)
                lines.append(f"| {row_line} |")
            
            return "\n".join(lines)
            
        except Exception as e:
            logger.error(f"Error normalizing table: {str(e)}")
            # Fallback to simple table text
            return self._table_to_text(table)
    
    def _generate_heuristic_table_summary(self, headers: List[str], rows: List[List[str]]) -> str:
        """
        Generate heuristic summary WITHOUT using LLM.
        
        Uses min/max/compare heuristics to summarize trends.
        Preserves units, year order, and exact numbers.
        
        Args:
            headers: Table headers
            rows: Cleaned data rows (numeric only)
            
        Returns:
            Short summary string
        """
        try:
            if not rows:
                return "No data available"
            
            # Find numeric columns
            numeric_columns = []
            for col_idx, header in enumerate(headers):
                # Check if this column has numeric values in all rows
                is_numeric = True
                for row in rows:
                    if col_idx < len(row):
                        cell = row[col_idx].strip()
                        if not re.search(r'[\d,]+\.?\d*', cell):
                            is_numeric = False
                            break
                if is_numeric:
                    numeric_columns.append((col_idx, header))
            
            if not numeric_columns:
                return "Table contains text data"
            
            # Generate summary for each numeric column
            summary_parts = []
            for col_idx, col_name in numeric_columns[:3]:  # Limit to first 3 numeric columns
                values = []
                for row in rows:
                    if col_idx < len(row):
                        cell = row[col_idx].strip()
                        # Extract numeric value
                        match = re.search(r'([\d,]+\.?\d*)', cell)
                        if match:
                            # Remove commas and convert to float
                            num_str = match.group(1).replace(',', '')
                            try:
                                value = float(num_str)
                                values.append(value)
                            except ValueError:
                                pass
                
                if values:
                    # Calculate min/max
                    min_val = min(values)
                    max_val = max(values)
                    
                    # Detect trend
                    if len(values) >= 2:
                        first_val = values[0]
                        last_val = values[-1]
                        change = last_val - first_val
                        change_pct = (change / first_val) * 100 if first_val != 0 else 0
                        
                        if change_pct > 5:
                            trend = f"increased from {min_val:,.0f} to {max_val:,.0f}"
                        elif change_pct < -5:
                            trend = f"decreased from {max_val:,.0f} to {min_val:,.0f}"
                        else:
                            trend = f"stable around {(min_val + max_val)/2:,.0f}"
                    else:
                        trend = f"ranges from {min_val:,.0f} to {max_val:,.0f}"
                    
                    summary_parts.append(f"{col_name}: {trend}")
            
            if summary_parts:
                return " ".join(summary_parts)
            else:
                return "Table data available"
                
        except Exception as e:
            logger.warning(f"Error generating heuristic summary: {str(e)}")
            return "Table data available"
    
    def _table_to_text(self, table: Dict[str, Any]) -> str:
        """
        Convert table to inline text representation (fallback).
        """
        headers = table.get('headers', [])
        rows = table.get('rows', [])
        
        lines = []
        lines.append("TABLE")
        
        if headers:
            lines.append(f"Columns: {', '.join(headers)}")
        
        for row in rows:
            if row:
                row_text = ' | '.join(cell for cell in row if cell)
                lines.append(row_text)
        
        return '\n'.join(lines)
    
    def _extract_with_camelot(self, file_path: str, page_num: int, flavor: str) -> List[Dict[str, Any]]:
        """Extract tables using Camelot."""
        tables = []
        try:
            tables_df = camelot.read_pdf(
                file_path,
                pages=str(page_num + 1),
                flavor=flavor,
                suppress_stdout=True
            )
            
            for table_idx, table in enumerate(tables_df):
                try:
                    df = table.df
                    if df.empty:
                        continue
                    
                    headers = [str(val).strip() for val in df.iloc[0].values if str(val) != 'nan']
                    rows = []
                    for _, row in df.iloc[1:].iterrows():
                        row_data = [str(val).strip() if str(val) != 'nan' else "" for val in row.values]
                        if any(row_data):
                            rows.append(row_data)
                    
                    if not headers or not rows:
                        continue
                    
                    table_data = {
                        'table_id': f"table_{table_idx}_page_{page_num + 1}",
                        'headers': headers,
                        'rows': rows,
                    }
                    
                    tables.append(table_data)
                    
                except Exception as e:
                    logger.warning(f"Error extracting Camelot table {table_idx}: {str(e)}")
                    continue
            
        except Exception as e:
            logger.debug(f"Camelot ({flavor}) failed: {str(e)}")
        
        return tables
    
    def _extract_with_tabula(self, file_path: str, page_num: int) -> List[Dict[str, Any]]:
        """Extract tables using Tabula."""
        tables = []
        try:
            tables_df = tabula.read_pdf(
                file_path,
                pages=page_num + 1,
                multiple_tables=True,
                silent=True
            )
            
            for table_idx, df in enumerate(tables_df):
                try:
                    if df.empty:
                        continue
                    
                    df.columns = [str(col).strip() for col in df.columns]
                    headers = list(df.columns)
                    
                    rows = []
                    for _, row in df.iterrows():
                        row_data = [str(val).strip() if str(val) != 'nan' else "" for val in row.values]
                        if any(row_data):
                            rows.append(row_data)
                    
                    if not headers or not rows:
                        continue
                    
                    table_data = {
                        'table_id': f"table_{table_idx}_page_{page_num + 1}",
                        'headers': headers,
                        'rows': rows,
                    }
                    
                    tables.append(table_data)
                    
                except Exception as e:
                    logger.warning(f"Error extracting Tabula table {table_idx}: {str(e)}")
                    continue
            
        except Exception as e:
            logger.debug(f"Tabula failed: {str(e)}")
        
        return tables
    
    def _extract_with_pymupdf(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract tables using PyMuPDF."""
        tables = []
        try:
            table_list = page.find_tables()
            
            for table_idx, table in enumerate(table_list):
                try:
                    headers = []
                    rows = []
                    
                    if table.rows:
                        headers = [cell.strip() if cell else "" for cell in table.rows[0]]
                        for row in table.rows[1:]:
                            row_data = [cell.strip() if cell else "" for cell in row]
                            if any(row_data):
                                rows.append(row_data)
                    
                    if not headers or not rows:
                        continue
                    
                    table_data = {
                        'table_id': f"table_{table_idx}_page_{page_num}",
                        'headers': headers,
                        'rows': rows,
                    }
                    
                    tables.append(table_data)
                    
                except Exception as e:
                    logger.warning(f"Error extracting PyMuPDF table {table_idx}: {str(e)}")
                    continue
            
        except Exception as e:
            logger.debug(f"PyMuPDF table extraction failed: {str(e)}")
        
        return tables
    
    def _extract_images_as_inline_text(self, page, page_num: int, doc_id: str, filename: str) -> Tuple[List[str], Dict[str, Any]]:
        """
        Extract images and generate inline text for RAG.
        
        Returns:
            Tuple of (image_texts list, image_metadata dict)
        """
        image_texts = []
        image_ids = []
        images_per_page_limit = 1
        images_processed = 0
        
        try:
            image_list = page.get_images()
            
            for img_idx, img in enumerate(image_list):
                if images_processed >= images_per_page_limit:
                    break
                
                try:
                    xref = img[0]
                    base_image = page.parent.extract_image(xref)
                    image_data = base_image["image"]
                    
                    # Check image size
                    if PIL_AVAILABLE:
                        image = Image.open(io.BytesIO(image_data))
                        width, height = image.size
                        
                        if width < self.MIN_IMAGE_SIZE or height < self.MIN_IMAGE_SIZE:
                            logger.info(f"Skipping small image {img_idx}: {width}x{height}")
                            continue
                    else:
                        width, height = 0, 0
                    
                    # Save image
                    image_id = f"{doc_id}_page_{page_num}_img_{img_idx}"
                    image_filename = f"{image_id}.png"
                    image_path = self.images_path / image_filename
                    
                    if PIL_AVAILABLE:
                        image.save(image_path, "PNG")
                    else:
                        with open(image_path, "wb") as img_file:
                            img_file.write(image_data)
                    
                    # Generate caption and description using LLM
                    inline_text = self._process_image_with_llm(image_data, page, page_num, image_id, doc_id)
                    
                    if inline_text:
                        image_texts.append(inline_text)
                        image_ids.append(image_id)
                        images_processed += 1
                        logger.info(f"Extracted image {image_id} from page {page_num}")
                    
                except Exception as e:
                    logger.warning(f"Error extracting image {img_idx}: {str(e)}")
                    continue
        
        except Exception as e:
            logger.debug(f"Image extraction failed for page {page_num}: {str(e)}")
        
        image_metadata = {
            'image_count': images_processed,
            'image_ids': image_ids
        }
        
        return image_texts, image_metadata
    
    def _process_image_with_llm(self, image_data: bytes, page, page_num: int, image_id: str, doc_id: str) -> Optional[str]:
        """
        Process image with LLM for caption and description generation.
        
        Implements:
        1. Check cache first (SHA256 hash)
        2. Detect existing caption in text
        3. Call LLM only if needed
        4. Retry logic with timeout
        5. Fallback on failure
        
        Args:
            image_data: Image bytes
            page: PyMuPDF page object
            page_num: Page number
            image_id: Unique image identifier
            doc_id: Document ID
            
        Returns:
            Inline text with format:
            [IMAGE]
            Caption: <caption>
            Description: <description>
        """
        try:
            # Calculate image hash for caching
            image_hash = self._get_image_hash(image_data)
            
            # Check cache first
            if image_hash in self.image_cache:
                cached_data = self.image_cache[image_hash]
                logger.info(f"Using cached caption for image {image_id}")
                caption = cached_data.get('caption', 'Image extracted from document')
                description = cached_data.get('description', '')
                
                # Format output
                if description:
                    return f"[IMAGE]\nCaption: {caption}\nDescription: {description}"
                else:
                    return f"[IMAGE]\nCaption: {caption}"
            
            # Not in cache - detect existing caption in text
            text_content = page.get_text()
            existing_caption = self._detect_caption_near_image(text_content)
            
            # Determine if we need to call LLM
            if existing_caption:
                # Caption exists - check if we need description
                # Only call LLM for description if image is large enough
                if PIL_AVAILABLE:
                    image = Image.open(io.BytesIO(image_data))
                    width, height = image.size
                    
                    if width >= self.MIN_IMAGE_SIZE_FOR_LLM and height >= self.MIN_IMAGE_SIZE_FOR_LLM:
                        # Call LLM for description only
                        description = self._call_llm_with_retry(image_data, description_only=True, existing_caption=existing_caption)
                    else:
                        description = ""
                else:
                    description = ""
                
                # Save to cache
                self.image_cache[image_hash] = {
                    'caption': existing_caption,
                    'description': description,
                    'timestamp': datetime.now().isoformat()
                }
                
                # Format output
                if description:
                    return f"[IMAGE]\nCaption: {existing_caption}\nDescription: {description}"
                else:
                    return f"[IMAGE]\nCaption: {existing_caption}"
            else:
                # No caption - call LLM for both caption and description
                if PIL_AVAILABLE:
                    image = Image.open(io.BytesIO(image_data))
                    width, height = image.size
                    
                    if width >= self.MIN_IMAGE_SIZE_FOR_LLM and height >= self.MIN_IMAGE_SIZE_FOR_LLM:
                        # Call LLM for caption and description
                        caption, description = self._call_llm_with_retry(image_data, description_only=False)
                    else:
                        # Fallback for small images
                        caption = "Image extracted from document"
                        description = ""
                else:
                    caption = "Image extracted from document"
                    description = ""
                
                # Save to cache
                self.image_cache[image_hash] = {
                    'caption': caption,
                    'description': description,
                    'timestamp': datetime.now().isoformat()
                }
                
                # Format output
                if description:
                    return f"[IMAGE]\nCaption: {caption}\nDescription: {description}"
                else:
                    return f"[IMAGE]\nCaption: {caption}"
                    
        except Exception as e:
            logger.warning(f"Error processing image {image_id}: {str(e)}")
            # Fallback on any error
            return None
    
    def _call_llm_with_retry(self, image_data: bytes, description_only: bool = False, existing_caption: str = "") -> Union[Tuple[str, str], str]:
        """
        Call OpenRouter Vision API with retry logic.
        
        FIXED: Uses correct vision model (gemini-1.5-flash) and proper payload format.
        Handles 400 errors gracefully with fallback.
        
        Args:
            image_data: Image bytes
            description_only: If True, only generate description
            existing_caption: Existing caption (for description-only mode)
            
        Returns:
            If description_only=False: (caption, description)
            If description_only=True: description only
        """
        for attempt in range(self.MAX_LLM_RETRIES + 1):
            try:
                base64_image = base64.b64encode(image_data).decode('utf-8')
                
                headers = {
                    "Authorization": f"Bearer {Config.OPENROUTER_API_KEY}",
                    "Content-Type": "application/json"
                }
                
                # Build system prompt (MANDATORY per requirements)
                system_prompt = "You are a professional financial document analyst. Describe images objectively for banking, AML, KYC, and compliance use. Focus on data, trends, charts, tables, and meaning. Avoid storytelling. Avoid speculation."
                
                # Build user prompt (MANDATORY per requirements)
                if description_only:
                    user_prompt = """Analyze this image from a financial document.
If it is a chart or graph:
- describe axes
- describe trend
- describe values if visible
- describe what it implies (increase, decrease, stability)
If it is a diagram or photo:
- describe what is shown
- describe its relevance to business or risk
Return JSON:
{
  "caption": "short factual caption",
  "description": "detailed professional description"
}"""
                else:
                    user_prompt = """Analyze this image from a financial document.
If it is a chart or graph:
- describe axes
- describe trend
- describe values if visible
- describe what it implies (increase, decrease, stability)
If it is a diagram or photo:
- describe what is shown
- describe its relevance to business or risk
Return JSON:
{
  "caption": "short factual caption",
  "description": "detailed professional description"
}"""
                
                # FIX: Use Gemini 1.5 Flash (vision-capable) instead of generic VISION_MODEL
                # Also ensure correct payload format with "url" field
                # CRITICAL: Add temperature and max_tokens for deterministic output
                data = {
                    "model": "google/gemini-1.5-flash:free",  # Vision-capable model
                    "temperature": 0.1,  # Low temperature for deterministic output
                    "max_tokens": 300,  # Limit response length
                    "messages": [
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": user_prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ]
                }
                
                logger.info(f"Calling LLM API (attempt {attempt + 1}/{self.MAX_LLM_RETRIES + 1}) for image")
                
                response = requests.post(
                    f"{Config.OPENROUTER_BASE_URL}/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=self.LLM_TIMEOUT
                )
                
                # FIX: Handle 400 errors specifically
                if response.status_code == 200:
                    result = response.json()
                    content = result["choices"][0]["message"]["content"]
                    
                    # Parse JSON response
                    try:
                        # Extract JSON from content (LLM might return extra text)
                        json_start = content.find('{')
                        json_end = content.rfind('}') + 1
                        
                        if json_start >= 0 and json_end > json_start:
                            json_content = content[json_start:json_end]
                            parsed = json.loads(json_content)
                            
                            caption = parsed.get('caption', 'Financial document image')
                            description = parsed.get('description', 'Financial chart or diagram from banking document')
                            
                            # Clean up caption
                            caption = caption.strip()
                            if len(caption) > 100:
                                caption = caption[:100]
                            
                            logger.info(f"✅ LLM generated caption: {caption[:50]}...")
                            
                            if description_only:
                                return description.strip()
                            else:
                                return caption, description.strip()
                        else:
                            # Failed to parse JSON - use meaningful fallback
                            logger.warning(f"Failed to parse LLM JSON response, using meaningful fallback")
                            if description_only:
                                return "Financial document image showing data or visual information"
                            else:
                                return "Financial document image", "Financial document image showing data or visual information"
                            
                    except json.JSONDecodeError as e:
                        logger.warning(f"Failed to parse LLM JSON: {str(e)}")
                        if description_only:
                            return "Financial document image showing data or visual information"
                        else:
                            return "Financial document image", "Financial document image showing data or visual information"
                elif response.status_code == 400:
                    # FIX: Log 400 error and retry aggressively
                    logger.error(f"❌ OpenRouter 400 error on attempt {attempt + 1}")
                    logger.error(f"Response: {response.text[:200]}")
                    if attempt < self.MAX_LLM_RETRIES:
                        logger.info(f"🔄 Retrying LLM call...")
                        continue
                    # Last attempt - use meaningful fallback
                    logger.warning(f"⚠️  All retries exhausted, using meaningful fallback")
                    if description_only:
                        return "Financial document image showing data or visual information from banking or AML compliance context"
                    else:
                        return "Financial document image", "Financial document image showing data or visual information from banking or AML compliance context"
                else:
                    logger.warning(f"OpenRouter Vision API error: {response.status_code}")
                    logger.warning(f"Response: {response.text[:200]}")
                    if attempt < self.MAX_LLM_RETRIES:
                        logger.info(f"🔄 Retrying LLM call...")
                        continue
                    
                    # Last attempt - use meaningful fallback
                    if description_only:
                        return "Financial document image showing data or visual information from banking or AML compliance context"
                    else:
                        return "Financial document image", "Financial document image showing data or visual information from banking or AML compliance context"
                        
            except requests.exceptions.Timeout:
                logger.warning(f"LLM API timeout (attempt {attempt + 1})")
                if attempt < self.MAX_LLM_RETRIES:
                    logger.info(f"Retrying LLM call...")
                    continue
                
                # Fallback - never crash pipeline
                if description_only:
                    return ""
                else:
                    return "Image extracted from document", ""
                    
            except Exception as e:
                logger.warning(f"LLM API error (attempt {attempt + 1}): {str(e)}")
                if attempt < self.MAX_LLM_RETRIES:
                    logger.info(f"Retrying LLM call...")
                    continue
                
                # Fallback - never crash pipeline
                if description_only:
                    return ""
                else:
                    return "Image extracted from document", ""
        
        # Should never reach here
        if description_only:
            return ""
        else:
            return "Image extracted from document", ""
    
    def _detect_caption_near_image(self, text_content: str) -> str:
        """Detect caption text."""
        caption_keywords = ['Figure', 'Fig.', 'Source', 'Note', 'Chart', 'Table']
        
        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if any(keyword.lower() in line.lower() for keyword in caption_keywords):
                if ':' in line:
                    caption = line.split(':', 1)[1].strip()
                    if len(caption) > 5:
                        return caption
        
        return ""
    
    def _process_docx(self, file_path: str, filename: str, doc_id: str) -> Dict[str, Any]:
        """Process DOCX file with RAG-first extraction."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing")
        
        document = {
            'doc_id': doc_id,
            'filename': filename,
            'pages': []
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract all text
            text_lines = []
            table_texts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_lines.append(para.text.strip())
            
            # Extract tables as inline text
            for table_idx, table in enumerate(doc.tables):
                table_text = self._docx_table_to_inline_text(table, table_idx)
                if table_text:
                    table_texts.append(table_text)
            
            # Merge all content
            all_content = text_lines + table_texts
            full_text = '\n\n'.join(all_content)
            
            # Create single page
            page_data = {
                'page_number': 1,
                'text': full_text.strip(),
                'metadata': {
                    'has_images': False,
                    'image_count': 0,
                    'image_ids': [],
                    'has_tables': len(table_texts) > 0,
                    'ocr_processed': False
                }
            }
            
            document['pages'] = [page_data]
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            raise Exception(f"DOCX processing failed: {str(e)}")
        
        return document
    
    def _docx_table_to_inline_text(self, table, table_idx: int) -> str:
        """Convert DOCX table to inline text."""
        if not table.rows or len(table.rows) < 2:
            return ""
        
        headers = [cell.text.strip() for cell in table.rows[0]]
        rows = [[cell.text.strip() for cell in row] for row in table.rows[1:]]
        
        lines = []
        lines.append("TABLE")
        
        if headers:
            lines.append(f"Columns: {', '.join(headers)}")
        
        for row in rows:
            if row:
                row_text = ' | '.join(cell for cell in row if cell)
                lines.append(row_text)
        
        return '\n'.join(lines)
    
    def _process_text_file(self, file_path: str, filename: str, doc_id: str, file_type: str) -> Dict[str, Any]:
        """Process text or markdown file with RAG-first extraction."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise Exception("Could not decode text file with any common encoding")
        
        # Clean and merge text
        clean_text = self._clean_and_merge_text(content)
        
        document = {
            'doc_id': doc_id,
            'filename': filename,
            'pages': [{
                'page_number': 1,
                'text': clean_text.strip(),
                'metadata': {
                    'has_images': False,
                    'image_count': 0,
                    'image_ids': [],
                    'has_tables': False,
                    'ocr_processed': False
                }
            }]
        }
        
        return document
    
    def _extract_text_with_ocr(self, file_path: str, page_num: int) -> str:
        """Extract text from PDF page using OCR."""
        try:
            images = pdf2image.convert_from_path(
                file_path,
                first_page=page_num + 1,
                last_page=page_num + 1,
                dpi=300
            )
            
            if images:
                text = pytesseract.image_to_string(images[0])
                logger.info(f"OCR extracted {len(text)} characters from page {page_num + 1}")
                return text
            
        except Exception as e:
            logger.error(f"OCR extraction failed for page {page_num + 1}: {str(e)}")
        
        return ""


# Backward compatibility alias
DataLoader = MultimodalDataLoader