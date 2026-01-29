"""
DOCX File Loader (Enhanced)
============================

WHAT IS THIS LOADER?
--------------------
This loader handles DOCX (Microsoft Word) files.
DOCX is the default format for Microsoft Word documents since 2007.

ENHANCEMENTS IN THIS VERSION:
-----------------------------
1. TEXT NORMALIZATION: Cleans extracted text for better quality
2. IMPROVED TABLE EXTRACTION: Better handling of complex tables
3. HEADER/FOOTER HANDLING: Optional extraction of headers/footers
4. PARAGRAPH PRESERVATION: Maintains document structure

DOCX FILE STRUCTURE:
--------------------
A DOCX file is actually a ZIP archive containing:
- XML files with document content
- Styles and formatting information
- Embedded images and media
- Document properties

The main content is in word/document.xml inside the ZIP.

WHAT LIBRARY DO WE USE?
-----------------------
We use python-docx to extract text from DOCX files.
python-docx is a popular library that:
- Opens DOCX files
- Provides access to paragraphs, tables, etc.
- Handles the complex XML structure internally

LIMITATIONS:
------------
- Only works with .docx files (not old .doc format)
- Images are not extracted (only text)
- Complex formatting may not be preserved
- Headers/footers require extra handling

For .doc files (old Word format), you would need:
- antiword (command-line tool)
- textract
- LibreOffice conversion
"""

import os
from typing import List, Optional

from src.ingestion.document_loader.base import DocumentLoader, LoadedDocument
from src.ingestion.text_utils import (
    TextNormalizer,
    NormalizationConfig,
    normalize_text
)


class DOCXLoader(DocumentLoader):
    """
    Document loader for DOCX (Microsoft Word) files with text normalization.

    This loader uses python-docx to extract text from Word documents.
    It extracts text from paragraphs and tables with optional normalization.

    Supported formats:
    - .docx (Word 2007 and later)

    NOT supported:
    - .doc (older Word format)

    Features:
    - Paragraph and table extraction
    - Text normalization (whitespace, control chars)
    - Document property extraction (title, author, etc.)
    - Optional header/footer extraction

    Requirements:
    - python-docx package: pip install python-docx

    Example:
        loader = DOCXLoader()
        doc = loader.load("report.docx")
        print(doc.text)
    """

    def __init__(
        self,
        normalize: bool = True,
        include_tables: bool = True,
        include_headers_footers: bool = False
    ):
        """
        Initialize the DOCX loader.

        Args:
            normalize: Whether to normalize extracted text (default: True).
                      Normalization removes control characters, normalizes
                      whitespace, and preserves paragraph boundaries.

            include_tables: Whether to extract text from tables (default: True).
                           Tables are added as tab-separated rows.

            include_headers_footers: Whether to include header/footer text
                                    (default: False). Headers/footers often
                                    contain repeated content that may not
                                    be useful for RAG.

        Example:
            # Default settings
            loader = DOCXLoader()

            # Include headers and footers
            loader = DOCXLoader(include_headers_footers=True)

            # Skip tables
            loader = DOCXLoader(include_tables=False)
        """
        self.normalize = normalize
        self.include_tables = include_tables
        self.include_headers_footers = include_headers_footers

        self.normalizer = TextNormalizer(NormalizationConfig(
            remove_control_chars=True,
            normalize_whitespace=True,
            normalize_newlines=True,
            remove_page_markers=False,
            detect_headers_footers=False,  # We handle this explicitly
            fix_ocr_artifacts=False,
            preserve_paragraphs=True,
            strip_lines=True
        ))

    @property
    def supported_extensions(self) -> List[str]:
        """
        Return supported file extensions.

        Note: We only support .docx, not .doc (old format)

        Returns:
            List containing ".docx"
        """
        return [".docx", ".DOCX"]

    def load(self, file_path: str) -> LoadedDocument:
        """
        Load a DOCX file and extract its text content.

        This method:
        1. Opens the DOCX file
        2. Extracts text from all paragraphs
        3. Extracts text from tables (if enabled)
        4. Optionally extracts headers/footers
        5. Normalizes the text (if enabled)
        6. Returns combined text with metadata

        Args:
            file_path: Path to the DOCX file.

        Returns:
            LoadedDocument with extracted text and metadata.

        Raises:
            FileNotFoundError: If file doesn't exist.
            ImportError: If python-docx is not installed.
            ValueError: If file cannot be read.
        """
        # =====================================================================
        # Step 1: Check if python-docx is installed
        # =====================================================================
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required to read DOCX files. "
                "Install it with: pip install python-docx"
            )

        # =====================================================================
        # Step 2: Validate the file exists
        # =====================================================================
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # =====================================================================
        # Step 3: Get file information
        # =====================================================================
        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)

        print(f"[DOCXLoader] Loading DOCX: {file_name}")
        print(f"[DOCXLoader] File size: {file_size} bytes")

        # =====================================================================
        # Step 4: Open and read the DOCX file
        # =====================================================================
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Could not read DOCX file: {e}")

        # =====================================================================
        # Step 5: Extract text from paragraphs
        # =====================================================================
        # Paragraphs are the main content containers in Word documents.
        # Each paragraph is a separate block of text.

        paragraph_texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraph_texts.append(text)

        print(f"[DOCXLoader] Found {len(paragraph_texts)} paragraphs")

        # =====================================================================
        # Step 6: Extract text from tables
        # =====================================================================
        table_texts = []
        table_count = 0

        if self.include_tables and doc.tables:
            for table_idx, table in enumerate(doc.tables, start=1):
                table_rows = []
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Clean cell text
                            cell_text = cell_text.replace('\n', ' ')
                            row_texts.append(cell_text)
                    if row_texts:
                        # Join cells with tabs (table-like format)
                        table_rows.append("\t".join(row_texts))

                if table_rows:
                    table_texts.extend(table_rows)
                    table_count += 1

            if table_count:
                print(f"[DOCXLoader] Found {table_count} tables")

        # =====================================================================
        # Step 7: Extract headers and footers (optional)
        # =====================================================================
        header_footer_texts = []

        if self.include_headers_footers:
            try:
                # Access document sections for headers/footers
                for section in doc.sections:
                    # Header
                    if section.header:
                        for para in section.header.paragraphs:
                            text = para.text.strip()
                            if text:
                                header_footer_texts.append(f"[Header] {text}")

                    # Footer
                    if section.footer:
                        for para in section.footer.paragraphs:
                            text = para.text.strip()
                            if text:
                                header_footer_texts.append(f"[Footer] {text}")

                if header_footer_texts:
                    print(f"[DOCXLoader] Found {len(header_footer_texts)} header/footer items")

            except Exception as e:
                print(f"[DOCXLoader] Warning: Could not extract headers/footers: {e}")

        # =====================================================================
        # Step 8: Combine all text
        # =====================================================================
        all_text_parts = []

        # Add headers/footers first (if any)
        if header_footer_texts:
            all_text_parts.append("--- Headers/Footers ---")
            all_text_parts.extend(header_footer_texts)
            all_text_parts.append("")  # Blank line separator

        # Add main content
        if paragraph_texts:
            all_text_parts.extend(paragraph_texts)

        # Add tables at the end
        if table_texts:
            all_text_parts.append("")  # Blank line separator
            all_text_parts.append("--- Tables ---")
            all_text_parts.extend(table_texts)

        full_text = "\n".join(all_text_parts)
        original_char_count = len(full_text)

        print(f"[DOCXLoader] Total characters (before normalization): {original_char_count}")

        # =====================================================================
        # Step 9: Normalize the text
        # =====================================================================
        if self.normalize:
            full_text = self.normalizer.normalize(full_text)
            print(f"[DOCXLoader] Normalized text ({original_char_count} → {len(full_text)} chars)")

        # =====================================================================
        # Step 10: Extract document properties (metadata)
        # =====================================================================
        doc_properties = {}
        try:
            core_props = doc.core_properties
            doc_properties = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
            # Remove empty values
            doc_properties = {k: v for k, v in doc_properties.items() if v}
        except Exception:
            # If we can't get properties, that's okay
            pass

        # =====================================================================
        # Step 11: Build metadata
        # =====================================================================
        metadata = {
            "source": file_name,
            "file_type": "docx",
            "file_path": file_path,
            "file_size_bytes": file_size,
            "paragraph_count": len(paragraph_texts),
            "table_count": table_count,
            "char_count": len(full_text),
            "original_char_count": original_char_count,
            "normalized": self.normalize,
            "includes_tables": self.include_tables,
            "includes_headers_footers": self.include_headers_footers,
            **doc_properties  # Include document properties
        }

        # =====================================================================
        # Step 12: Return the loaded document
        # =====================================================================
        return LoadedDocument(text=full_text, metadata=metadata)
